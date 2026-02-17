"""DOCX parser — extracts text and tables from Word documents."""
import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.table import Table

from parsers import DocumentDict

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parse DOCX files into the standard document dict format."""

    def parse(self, path: str) -> DocumentDict:
        doc = DocxDocument(path)
        pages = []
        tables = []

        # Extract paragraphs (DOCX doesn't have pages, treat as single page)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        title = Path(path).stem
        if full_text:
            for line in full_text:
                if len(line.strip()) > 10:
                    title = line.strip()
                    break
            body = "\n\n".join(full_text)
            markdown = f"# {title}\n\n{body}"
            pages.append({"text": markdown, "page": 1})

        # Extract tables
        for i, table in enumerate(doc.tables):
            table_data = self._extract_table(table)
            if table_data:
                tables.append({"data": table_data, "index": i})

        return {
            "pages": pages,
            "source": Path(path).name,
            "title": title,
            "creation_date": "",
            "authors": "",
            "tables": tables,
        }

    @staticmethod
    def _extract_table(table: Table) -> list[list[str]]:
        """Extract a table as a list of rows (each row is a list of cell texts)."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        return rows
