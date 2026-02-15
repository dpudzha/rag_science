"""Tests for parsers/docx_parser.py."""
import pytest
from pathlib import Path
from docx import Document as DocxDocument


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample DOCX file."""
    path = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("This is the title of the document")
    doc.add_paragraph("This is the introduction paragraph with enough text.")
    doc.add_paragraph("This section discusses methodology.")

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Method"
    table.cell(0, 1).text = "Accuracy"
    table.cell(1, 0).text = "CNN"
    table.cell(1, 1).text = "95.2"
    table.cell(2, 0).text = "RNN"
    table.cell(2, 1).text = "92.1"

    doc.save(str(path))
    return str(path)


class TestDOCXParser:
    def test_extracts_text(self, sample_docx):
        from parsers.docx_parser import DOCXParser
        parser = DOCXParser()
        result = parser.parse(sample_docx)
        assert result["pages"]
        assert "introduction" in result["pages"][0]["text"].lower()

    def test_extracts_title(self, sample_docx):
        from parsers.docx_parser import DOCXParser
        parser = DOCXParser()
        result = parser.parse(sample_docx)
        assert result["title"]
        assert len(result["title"]) > 10

    def test_extracts_tables(self, sample_docx):
        from parsers.docx_parser import DOCXParser
        parser = DOCXParser()
        result = parser.parse(sample_docx)
        assert result["tables"]
        table = result["tables"][0]
        assert "Method" in table["data"][0]
        assert "CNN" in table["data"][1]

    def test_source_is_filename(self, sample_docx):
        from parsers.docx_parser import DOCXParser
        parser = DOCXParser()
        result = parser.parse(sample_docx)
        assert result["source"] == "test.docx"

    def test_has_required_keys(self, sample_docx):
        from parsers.docx_parser import DOCXParser
        parser = DOCXParser()
        result = parser.parse(sample_docx)
        for key in ("pages", "source", "title", "creation_date", "authors"):
            assert key in result
